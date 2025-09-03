from typing import Dict, Any, List
from docflow.adapters.base import DocumentAdapter
from docx import Document
import re
import io
from pathlib import Path
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.text.run import Run

IMAGE_RE = re.compile(r"{{image:([^}]+)}}")
VAR_RE = re.compile(r"{{\s*([^}]+)\s*}}")

def clean_markdown_text(text: str) -> str:
    """Remove Markdown formatting and return clean text"""
    # Remove bold (**text**)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Remove italic (*text*)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # Convert bullet points to simple dashes
    text = re.sub(r'^[-\*]\s+', '- ', text, flags=re.MULTILINE)
    return text


class DocxAdapter(DocumentAdapter):
    def __init__(self, path: str):
        super().__init__(Path(path))
        self.doc = None

    def load(self):
        # mitigate occasional PermissionError on network/OneDrive paths by simple retry
        attempts = 0
        last_exc = None
        while attempts < 3:
            try:
                self.doc = Document(self.path)
                return
            except Exception as e:
                last_exc = e
                import time as _t
                _t.sleep(0.2 * (attempts + 1))
                attempts += 1
        if last_exc:
            raise last_exc

    def list_placeholders(self) -> List[str]:
        if self.doc is None:
            self.load()
        found = set()
        for para in self.doc.paragraphs:
            for run in para.runs:
                text = run.text or ''
                for m in VAR_RE.finditer(text):
                    name = m.group(1)
                    if ':' in name:
                        name = name.split(':', 1)[1]
                    found.add(name)
        return sorted(found)

    def apply(self, mapping: Dict[str, Any], global_vars: Dict[str, Any]):
        if self.doc is None:
            self.load()
        for para in self.doc.paragraphs:
            for run in para.runs:
                text = run.text
                if not text:
                    continue
                m = IMAGE_RE.search(text)
                if m:
                    key = m.group(1)
                    val = mapping.get(key) or global_vars.get(key)
                    if val:
                        # if bytes or path, insert image
                        if isinstance(val, (bytes, bytearray)):
                            bio = io.BytesIO(val)
                            def repl(mo):
                                nm = mo.group(1)
                                if ':' in nm:
                                    nm = nm.split(':',1)[1]
                                return str(mapping.get(nm, global_vars.get(nm, mo.group(0))))
                            run.text = VAR_RE.sub(repl, text)
                            para.add_run().add_picture(bio, width=Inches(4))
                        elif isinstance(val, str):
                            pth = Path(val)
                            if pth.exists():
                                def repl(mo):
                                    nm = mo.group(1)
                                    if ':' in nm:
                                        nm = nm.split(':',1)[1]
                                    return ''  # remove placeholder text, image will be inserted
                                run.text = VAR_RE.sub(repl, text)
                                try:
                                    para.add_run().add_picture(str(pth), width=Inches(4))
                                except Exception:
                                    # invalid/corrupt image - leave placeholder removed but continue
                                    pass
                            else:
                                # fallback: just replace other vars but keep text
                                def repl(mo):
                                    nm = mo.group(1)
                                    if ':' in nm:
                                        nm = nm.split(':',1)[1]
                                    return str(mapping.get(nm, global_vars.get(nm, mo.group(0))))
                                run.text = VAR_RE.sub(repl, text)
                        continue
                def repl_default(mo):
                    nm = mo.group(1)
                    if ':' in nm:
                        nm = nm.split(':',1)[1]
                    value = str(mapping.get(nm, global_vars.get(nm, mo.group(0))))
                    return value
                new_text = VAR_RE.sub(repl_default, text)
                if new_text != text:
                    # Check if the new text contains Markdown and clean it
                    if any(marker in new_text for marker in ['**', '*', '- ', '* ']):
                        new_text = clean_markdown_text(new_text)
                    run.text = new_text

    def save(self, out_path: str):
        if self.doc is None:
            self.load()
        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(out_path)

