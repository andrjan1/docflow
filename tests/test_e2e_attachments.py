import yaml
from pathlib import Path
from docflow.runtime.orchestrator import run_config


def test_e2e_run_with_kb_unified(tmp_path: Path):
    """Test the new unified KB system (formerly attachments)"""
    # create KB files
    kbdir = tmp_path / 'kb'
    kbdir.mkdir()
    (kbdir / 'note.md').write_text('# Note\nThis is a KB note about sales and installation')

    # create a simple docx template
    from docx import Document

    templates = tmp_path / 'templates'
    templates.mkdir()
    doc = Document()
    doc.add_paragraph('Greeting: {{greeting}}')
    doc.save(templates / 'demo_template.docx')

    # write config - UPDATED to use unified kb system
    cfg = {
        'project': {'base_dir': str(tmp_path), 'output_dir': 'output'},
        'ai': {'provider': 'mock'},
        'workflow': {
            'actions': [
                {
                    'id': 'gen1',
                    'type': 'generative',
                    'returns': 'image',
                    'prompt': 'Say hello to {{name}}',
                    'kb': {
                        'enabled': True,
                        'paths': [str(kbdir / '*.md')], 
                        'strategy': 'hybrid',  # Both text and upload
                        'as_text': True, 
                        'upload': False,
                        'max_chars': 10000
                    },
                }
            ],
            'templates': [{'path': str(templates / 'demo_template.docx'), 'adapter': 'docx'}],
        },
    }
    cfg_path = tmp_path / 'cfg.yaml'
    cfg_path.write_text(yaml.safe_dump(cfg))

    # run orchestrator
    out = run_config(str(cfg_path))
    # should have created an output file
    assert any(p.endswith('.docx') for p in out)


def test_kb_upload_strategy(tmp_path: Path):
    """Test the new upload strategy in unified KB system"""
    kbdir = tmp_path / 'kb'
    kbdir.mkdir()
    (kbdir / 'document.txt').write_text('Sample document content for upload testing')
    
    templates = tmp_path / 'templates'
    templates.mkdir()
    from docx import Document
    doc = Document()
    doc.add_paragraph('Analysis: {{analysis}}')
    doc.save(templates / 'report.docx')

    cfg = {
        'project': {'base_dir': str(tmp_path), 'output_dir': 'output'},
        'ai': {'provider': 'mock'},
        'workflow': {
            'actions': [
                {
                    'id': 'analyze_docs',
                    'type': 'generative',
                    'returns': 'text',
                    'prompt': 'Analyze the uploaded documents',
                    'kb': {
                        'enabled': True,
                        'paths': [str(kbdir / '*.txt')],
                        'strategy': 'upload',  # Upload files to AI provider
                        'upload': True,
                        'as_text': False,
                        'mime_type': 'text/plain'
                    },
                }
            ],
            'templates': [{'path': str(templates / 'report.docx'), 'adapter': 'docx'}],
        },
    }
    
    cfg_path = tmp_path / 'cfg.yaml'
    cfg_path.write_text(yaml.safe_dump(cfg))

    out = run_config(str(cfg_path))
    assert any(p.endswith('.docx') for p in out)


def test_kb_inline_strategy(tmp_path: Path):
    """Test the inline strategy (text extraction)"""
    kbdir = tmp_path / 'kb'
    kbdir.mkdir()
    (kbdir / 'info.md').write_text('# Information\nThis document contains important information.')
    
    templates = tmp_path / 'templates'
    templates.mkdir()
    from docx import Document
    doc = Document()
    doc.add_paragraph('Summary: {{summary}}')
    doc.save(templates / 'summary.docx')

    cfg = {
        'project': {'base_dir': str(tmp_path), 'output_dir': 'output'},
        'ai': {'provider': 'mock'},
        'workflow': {
            'actions': [
                {
                    'id': 'summarize',
                    'type': 'generative',
                    'returns': 'text',
                    'prompt': 'Summarize the following information: {{kb_text}}',
                    'kb': {
                        'enabled': True,
                        'paths': [str(kbdir / '*.md')],
                        'strategy': 'inline',  # Extract text and include in prompt
                        'as_text': True,
                        'max_chars': 5000
                    },
                }
            ],
            'templates': [{'path': str(templates / 'summary.docx'), 'adapter': 'docx'}],
        },
    }
    
    cfg_path = tmp_path / 'cfg.yaml'
    cfg_path.write_text(yaml.safe_dump(cfg))

    out = run_config(str(cfg_path))
    assert any(p.endswith('.docx') for p in out)
